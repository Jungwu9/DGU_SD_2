# -*- coding: utf-8 -*-
import os
import re
import json
import glob
import zipfile
import io
from io import BytesIO
from typing import List, Tuple, Optional
import time

import streamlit as st
import pdfplumber
from PIL import Image

import torch
from sentence_transformers import SentenceTransformer

# LangChain & Gemini
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.documents import Document
from langchain_core.runnables import RunnablePassthrough, RunnableLambda
from langchain_core.output_parsers import StrOutputParser
from langchain_google_genai import ChatGoogleGenerativeAI

# Word export
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    DocxDocument = None

# ==========================================
# [IMAGE BLOCK - 활성화됨]
# ==========================================
try:
    from google import genai
    from google.genai import types
except ImportError:
    genai = None
    types = None
# ==========================================
# [IMAGE BLOCK END]
# ==========================================


# ==========================================
# ==== 1. Configuration ====================
# ==========================================
DEFAULT_API_KEY = ""
FIXED_MODEL_PRO = "gemini-2.5-pro"    # 제안서 작성용 (고도화)
FIXED_MODEL_FLASH = "gemini-2.5-flash" # 요약용 (빠른 속도)
FIXED_MODEL_IMAGE = "gemini-3.1-flash-image-preview" # 이미지 생성용
SBERT_MODEL_FOR_MATCH = "sentence-transformers/paraphrase-multilingual-mpnet-base-v2"

# ==========================================
# ==== 2. Session State 초기화 (UI 리셋 방지) =
# ==========================================
if "summary_result" not in st.session_state:
    st.session_state.summary_result = ""
if "proposal_result" not in st.session_state:
    st.session_state.proposal_result = ""
if "docx_bytes" not in st.session_state:
    st.session_state.docx_bytes = None
if "template_guide" not in st.session_state:
    st.session_state.template_guide = ""
if "research_ev" not in st.session_state:
    st.session_state.research_ev = ""
if "goal_img" not in st.session_state:
    st.session_state.goal_img = None
if "fw_img" not in st.session_state:
    st.session_state.fw_img = None

# ==========================================
# ==== 3. Styles (가독성 높은 전문 테마) =====
# ==========================================
st.set_page_config(
    page_title="사업계획서 자동 생성기",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Noto+Serif+KR:wght@400;500;700&display=swap');

/* 차분하고 전문적인 배경화면 설정 (#f5f7f9) */
body, .stApp {
    background-color: #f5f7f9;
    color: #333333;
    font-family: "Noto Serif KR", "Batang", "바탕", serif !important;
}

/* 제목 영역 스타일 (로고 제거, 제목만 중앙 또는 좌측 정렬) */
.header-container {
    border-bottom: 2px solid #ea5414; /* 동국대 오렌지색 포인트 */
    padding-bottom: 15px;
    margin-bottom: 30px;
    background-color: transparent;
}
.header-title {
    color: #111111;
    font-size: 34px;
    font-weight: 700;
    margin: 0;
    font-family: "Noto Serif KR", "Batang", "바탕", serif !important;
}

/* 내용을 담는 하얀색 박스 스타일 */
.white-box {
    background-color: #ffffff;
    color: #111111;
    padding: 35px 45px;
    border-radius: 12px;
    box-shadow: 0 5px 15px rgba(0,0,0,0.05); /* 은은한 그림자 */
    margin-bottom: 25px;
    border: 1px solid #e1e4e8;
}

/* 화이트 박스 내부의 마크다운 텍스트 스타일 */
.white-box p, .white-box li, .white-box div, .white-box td, .white-box th {
    font-family: "Noto Serif KR", "Batang", "바탕", serif !important;
    line-height: 1.9 !important;
    font-size: 16px !important;
    color: #333333 !important;
}
.white-box h1, .white-box h2, .white-box h3, .white-box h4 {
    color: #111111 !important;
    font-family: "Noto Serif KR", "Batang", "바탕", serif !important;
    font-weight: 700;
}
.white-box h1 { border-bottom: 2px solid #ea5414; padding-bottom: 8px; margin-top: 25px; margin-bottom: 15px; font-size: 26px !important;}
.white-box h2 { font-size: 22px !important; margin-top: 20px;}
.white-box table { width: 100%; border-collapse: collapse; margin-bottom: 20px; border: 1px solid #cccccc;}
.white-box th, .white-box td { border: 1px solid #cccccc; padding: 10px; }
.white-box th { background-color: #f7f9fc; }

/* 이미지 플레이스홀더 및 컨테이너 박스 */
.image-placeholder {
    background-color: #f8f9fa;
    border: 2px dashed #cccccc;
    border-radius: 10px;
    padding: 50px 25px;
    text-align: center;
    color: #777777;
    margin-bottom: 25px;
}
.image-placeholder h4 { color: #555555 !important; }

/* 생성 버튼 흰색 스타일 설정 */
div.stButton > button:first-child {
    background-color: #ffffff !important;
    color: #ea5414 !important;
    border: 2px solid #ea5414 !important;
    font-weight: 700 !important;
    border-radius: 8px !important;
    padding: 12px 30px !important;
    font-size: 18px !important;
    transition: all 0.3s ease;
}
div.stButton > button:first-child:hover {
    background-color: #ea5414 !important;
    color: #ffffff !important;
    box-shadow: 0 4px 10px rgba(234, 84, 20, 0.3);
}

/* 사이드바 스타일 수정 */
[data-testid="stSidebar"] {
    background-color: #ffffff;
    border-right: 1px solid #e1e4e8;
}
[data-testid="stSidebar"] h1, [data-testid="stSidebar"] h2 {
    color: #111111;
}

/* 업로드 섹션 라벨 스타일 */
.stFileUploader label {
    color: #111111 !important;
    font-weight: 600 !important;
    font-size: 16px !important;
}
</style>
""", unsafe_allow_html=True)


# ==========================================
# ==== 4. Utilities ========================
# ==========================================
def normalize_space(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip()

def truncate_text(text: str, max_chars: int) -> str:
    text = text or ""
    return text[:max_chars] if len(text) > max_chars else text

def build_text_llm(api_key: str, model_name: str, temperature: float = 0.0):
    return ChatGoogleGenerativeAI(
        model=model_name,
        temperature=temperature,
        google_api_key=api_key
    )

def build_genai_client(api_key: str):
    if genai is None or types is None:
        st.error("❌ google-genai SDK가 설치되어 있지 않습니다. pip install google-genai pillow")
        st.stop()
    return genai.Client(api_key=api_key)

def extract_text_from_uploaded_pdf(uploaded_file, max_pages: Optional[int] = None) -> Tuple[str, str]:
    if uploaded_file is None:
        return "", ""
    text_parts = []
    try:
        pdf_bytes = uploaded_file.getvalue()
        with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
            pages = pdf.pages[:max_pages] if max_pages else pdf.pages
            for page in pages:
                t = page.extract_text() or ""
                if t.strip():
                    text_parts.append(t)
        return "\n".join(text_parts).strip(), uploaded_file.name
    except Exception as e:
        return f"(PDF 추출 실패: {e})", uploaded_file.name

def extract_docs_from_uploaded_pdf(uploaded_file) -> List[Document]:
    if uploaded_file is None:
        return []
    try:
        pdf_bytes = uploaded_file.getvalue()
        with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        if not text.strip():
            return []
        return [Document(page_content=text)]
    except Exception as e:
        st.error(f"공고문 PDF 읽기 오류: {e}")
        return []


# ==========================================
# ==== 5. Image Generation Functions =======
# ==========================================
def _parse_image_response(resp):
    out_text = []
    out_images = []
    try:
        if resp and resp.candidates:
            parts = resp.candidates[0].content.parts
            for part in parts:
                if getattr(part, "text", None):
                    out_text.append(part.text)
                if getattr(part, "inline_data", None) and getattr(part.inline_data, "data", None):
                    img = Image.open(BytesIO(part.inline_data.data)).convert("RGB")
                    out_images.append(img)
    except Exception:
        pass
    return ("\n".join(out_text).strip(), out_images[0] if out_images else None)

def generate_goal_image_with_genai(client, project_summary: str, template_guidance: str, research_evidence: str):
    prompt = f"""
    국가 R&D 사업계획서 삽입용 '연구 목적·비전·전략·기대효과' 개념도 1장을 생성할 것.

    절대 넣지 말 것:
    - 교수명, 대학/기관명, 긴 문장, 논문 제목, Figure 번호, 하단 캡션

    허용:
    - 짧은 박스 라벨, 화살표, 핵심 키워드

    필수 구성:
    - 문제 인식, 목표, 추진 전략 3~4개, 기대효과 등

    스타일:
    - 흰 배경, 깔끔한 인포그래픽, 사업계획서 삽입용, 단정한 한국어 라벨

    [사업 요약]
    {project_summary}
    [양식/공고 반영 가이드]
    {template_guidance}
    [논문 기반 기술요소]
    {research_evidence}
    """.strip()

    try:
        resp = client.models.generate_content(
            model=FIXED_MODEL_IMAGE,
            contents=[prompt],
            config=types.GenerateContentConfig(
                response_modalities=[types.Modality.IMAGE, types.Modality.TEXT],
                temperature=0.2,
            ),
        )
        return _parse_image_response(resp)
    except Exception as e:
        return (f"(목표 이미지 생성 실패: {e})", None)

def generate_framework_image_with_genai(client, project_summary: str, template_guidance: str, research_evidence: str):
    prompt = f"""
    국가 R&D 사업계획서 삽입용 '연구 프레임워크·추진체계' 도식 1장을 생성할 것.

    절대 넣지 말 것:
    - 교수명, 대학/기관명, 긴 문장, 논문 제목, Figure 번호, 큰 제목

    허용:
    - 짧은 모듈명, 화살표, 단계 구분, 입력/분석/평가/확산 흐름

    필수 구성:
    - 연구 자원/데이터/기반기술
    - 분석 체계 또는 방법론 묶음
    - 운영 및 지원 체계
    - 평가 및 성과 확산 흐름

    스타일:
    - 흰 배경, 깔끔한 인포그래픽, 단정한 한국어 라벨

    [사업 요약]
    {project_summary}
    [양식/공고 반영 가이드]
    {template_guidance}
    [논문 기반 기술요소]
    {research_evidence}
    """.strip()

    try:
        resp = client.models.generate_content(
            model=FIXED_MODEL_IMAGE,
            contents=[prompt],
            config=types.GenerateContentConfig(
                response_modalities=[types.Modality.IMAGE, types.Modality.TEXT],
                temperature=0.2,
            ),
        )
        return _parse_image_response(resp)
    except Exception as e:
        return (f"(프레임워크 이미지 생성 실패: {e})", None)


# ==========================================
# ==== 6. Prompt Definitions ===============
# ==========================================
def get_summary_prompt_template():
    return r"""
## 지시사항 (Instruction)
당신은 국가 연구개발 과제 공고문을 분석하여 핵심 정보를 추출하는 전문 연구 분석가입니다.
주어진 공고문 텍스트를 검토한 뒤, 아래 요약 양식에 따라 정확하고 상세하게 작성하세요.
추정이나 유추는 금지하며, 공고문에 기반하여 작성하세요.

---
## 입력 데이터 (RFP 원문)
{context}

---
## 사용자 질문
{input}

---
## 출력 양식 (요약 결과)

[사업 공고 핵심 요약]

### 1. 과제 목표
- ...

### 2. 연구 기간
- ...

### 3. 과제 예산
- ...

### 4. 지원 자격 및 형태
- ...

### 5. 사업 내용
- ...

### 6. 주요 평가 항목/중점 사항
- ...
"""

def get_template_analysis_prompt():
    return """
당신은 정부·재단 R&D 사업 사업계획서 양식 해석 전문가임.
입력으로 제공되는 문서는 다음 두 종류임.
  1) 사업 공고문
  2) 사업계획서 공식 양식(Template)

▶ 목표
공고·양식에서 실제 집필에 필요한 규칙만 뽑아서 "제안서 집필 가이드"를 생성하되,
특히 '양식(Template)'의 목차별 작성 의도와 요구사항, 힌트 텍스트를 매우 세밀하게 추출하여,
초안 작성자가 양식의 원래 의도에 완벽하게 맞추어 빈칸을 채우듯 작성할 수 있도록 지침을 구체화할 것.

▶ 양식 표(Table) 구조 분석 지침 (★ 최우선)
- 양식에 등장하는 모든 표(table)를 빠짐없이 열거할 것
- 각 표에 대해:
  1) 표 제목/위치(어느 목차에 속하는지)
  2) 표의 열(column) 헤더 목록
  3) 표의 행(row) 헤더 목록  
  4) 각 셀에 들어갈 내용의 성격(수치/서술문/항목명 등)
  5) 필수 기재 여부
- 표 내 빈칸, 체크박스, 기입란을 모두 채워야 하는 대상으로 간주할 것

▶ 음슴체 규칙 (최우선 적용)
  - 모든 서술문은 반드시 아래 어미로 종결할 것:
      ~임 / ~함 / ~됨 / ~있음 / ~없음 / ~필요함
      ~예정임 / ~기대됨 / ~추진함 / ~수행함 / ~확보함
      ~설치함 / ~운영함 / ~지원함 / ~마련함 / ~구성함
  - 절대 사용 금지 어미:
      ~합니다 / ~입니다 / ~습니다 / ~겠습니다
      ~이다 / ~한다 / ~된다 / ~한다고 / ~있다 / ~없다

▶ 출력 형식

# 제안서 집필 가이드

## 1. 문체·화자 규칙
## 2. 공식 목차 구조 및 섹션별 세부 작성 의도 (양식 밀착)
## 3. 양식 표(Table) 전체 목록 및 구조 분석 ← 핵심 섹션
## 4. 그림·도표 허용 위치 및 규칙
## 5. 외부 절차성 정보 (본문 사용 금지 목록)
## 6. 공식 인력 용어 정리
## 7. 표 셀별 기재 기준 (채워야 할 항목 체크리스트)

[공고문 텍스트]
{notice_text}

[양식 텍스트]
{template_text}
""".strip()

def get_research_grounding_prompt():
    return """
당신은 연구제안서용 '기술요소 근거 추출 전담 분석가'임.
입력은 동국대학교 소속 교수의 직접 업로드된 논문 PDF 발췌임.

목표:
- 동국대학교 연구팀의 독보적 역량으로 활용될 수 있도록 논문의 핵심 기술/학술요소를 매우 풍부하게 추출할 것
- 사업계획서 문안에 즉시 재사용 가능한 형태로 구체적이고 전문적으로 정리할 것
- 확인되지 않은 데이터 발명 금지

출력 형식:
## 확인된 연구/기술 주제 (동국대학교 연구진 기반)
## 동국대학교 연구진의 핵심 방법론·분석틀·운영개념
## 확인된 데이터·대상·연구환경
## 사업계획서 내 적극 반영 가능 포인트 (매우 구체적으로)

[논문 발췌]
{research_content}
""".strip()

def get_proposal_prompt():
    return """
당신은 대한민국 최고 수준의 R&D 사업계획서 작성 전문가임.
현재 '동국대학교 교수님'들을 대신하여, 업로드된 선행 논문을 바탕으로 실제 제출용 사업계획서 초안을 작성 중임.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━
[사업 공고 핵심 요약]
{project_summary}

[양식·공고 분석 결과 — 집필 가이드 (표 구조 포함)]
{template_guidance}

[논문 기반 기술요소]
{research_evidence}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━

━━━━ 1. 동국대학교 및 논문 역량 집중 부각 ━━━━
1. 본 사업의 주관기관은 반드시 "동국대학교 산학협력단" 또는 "동국대학교"로 명시할 것.
2. 제안서 양식의 각 목차 목적에 맞추어, 첨부된 [논문 기반 기술요소]를 아낌없이 활용해 아주 깊이 있고 방대하게 꽉 채워넣을 것. 

━━━━ 2. 양식(Template) 100% 밀착 작성 (빈칸 채우기 방식) ━━━━
1. [집필 가이드]에 명시된 '섹션별 요구사항과 작성 지침'에 정확히 들어맞는 내용만으로 "양식의 빈칸을 완벽하게 채워 넣듯" 철저히 양식을 따라갈 것.
2. 임의로 양식에 없는 목차 추가 절대 금지.

━━━━ 3. 표(Table) 작성 핵심 규칙 (★ 최우선) ━━━━
1. 양식에 있는 모든 표를 마크다운 표(|---|) 형식으로 재현하고 빠짐없이 채울 것.
2. 각 셀은 반드시 실제 서술 내용으로 채울 것 — 빈칸, "○", "〔입력 필요〕" 등 placeholder 절대 금지. 합리적 수치를 스스로 추정하여 기재할 것.

━━━━ 4. 필수 문체 — 음슴체 강제 ━━━━
모든 서술 문장 종결을 개조식 및 음슴체로 통일할 것.
  - 올바른 어미: ~임 / ~함 / ~됨 / ~있음 / ~없음 / ~필요함 / ~예정임
  - 절대 금지 어미: ~합니다 / ~입니다 / ~한다 / ~된다 / ~이다

━━━━ 5. 기타 제약 사항 ━━━━
1. 외부 선정평가·접수일정 등 절차성 정보 제외.
2. [[IMAGE:...]], [[TABLE:...]] 외 임의 태그 생성 절대 금지 (허용: [[IMAGE:GOAL]], [[IMAGE:FRAMEWORK]] 2개만).

━━━━ 출력 형식 ━━━━
집필 가이드의 공식 목차 구조를 그대로 따르되, 아래 구조로 시작하여 마지막 목차까지 생략 없이 작성할 것:

# [사업계획서 요약표 작성용 서술문]
## 설립·추진 목적
## 사업단(연구소) 사업 요약
## 연구인력 양성 및 기대효과

[매우 중요 경고] 
이후 공식 목차 순서대로 **마지막 목차까지 단 하나도 생략 없이** 작성할 것.
특히 [집필 가이드]에 명시된 **모든 표(Table)는 본문 내 해당 위치에 반드시 그려 넣고, 빈칸 없이 꽉 채울 것.**
""".strip()

def get_compliance_revision_prompt():
    return """
당신은 국가 R&D 사업계획서 품질관리 전문가임.
입력된 초안을 아래 규칙에 맞게 전체 검토·재작성할 것.

━━━━ 검토 항목 ━━━━
[1] 동국대학교 명시: 주관기관이 "동국대학교"로 명확히 기재되도록 할 것.
[2] 음슴체 전면 적용: ~합니다, ~입니다, ~한다 금지. ~함, ~임, ~됨으로 통일.
[3] 표(Table) 완성도 검토: "입력 필요", "-", 빈칸이 있는 셀은 합리적인 실제 내용과 수치로 채울 것.
[4] 금지 태그 제거: [[IMAGE:GOAL]], [[IMAGE:FRAMEWORK]] 외 임의의 [[IMAGE:...]], [[TABLE:...]] 태그 전면 삭제.
[5] 줄바꿈 기호 처리: 본문 내에 있는 <br> 또는 &lt;br&gt; 기호는 마크다운에서 인식 가능한 실제 엔터(줄바꿈)로 변환하여 출력할 것. 텍스트에 <br> 글자가 그대로 노출되지 않게 할 것.

━━━━ 입력 ━━━━
[집필 가이드]
{template_guidance}

[초안]
{draft}

━━━━ 출력 ━━━━
검토 완료된 사업계획서 전문을 그대로 출력할 것. 메타 서술 절대 금지.
""".strip()


# ==========================================
# ==== 7. LLM Orchestration ================
# ==========================================
def run_summarization(docs: List[Document], api_key: str) -> str:
    splitter = RecursiveCharacterTextSplitter(chunk_size=3000, chunk_overlap=500)
    texts = splitter.split_documents(docs)
    
    hf_emb = HuggingFaceEmbeddings(
        model_name=SBERT_MODEL_FOR_MATCH,
        model_kwargs={"device": "cuda" if torch.cuda.is_available() else "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )
    vector_store = FAISS.from_documents(texts, hf_emb)
    retriever = vector_store.as_retriever(search_kwargs={"k": 7})

    llm = build_text_llm(api_key, FIXED_MODEL_FLASH)
    prompt = ChatPromptTemplate.from_template(get_summary_prompt_template())
    
    chain = (
        {
            "context": retriever | RunnableLambda(lambda x: "\n\n".join(d.page_content for d in x)),
            "input": RunnablePassthrough(),
        }
        | prompt
        | llm
        | StrOutputParser()
    )
    question = "이 공고문의 내용을 '출력 양식'에 맞춰 정확하고 간결하게 요약해줘."
    summary_raw = chain.invoke(question)
    return summary_raw.strip()

def analyze_notice_and_template(llm, notice_text: str, template_text: str):
    chain = ChatPromptTemplate.from_template(get_template_analysis_prompt()) | llm | StrOutputParser()
    return chain.invoke({
        "notice_text": truncate_text(notice_text, 18000) if notice_text else "(공고문 미제공)",
        "template_text": truncate_text(template_text, 50000) if template_text else "(양식 미제공)",
    })

def build_research_evidence(llm, research_content: str):
    chain = ChatPromptTemplate.from_template(get_research_grounding_prompt()) | llm | StrOutputParser()
    return chain.invoke({
        "research_content": truncate_text(research_content, 30000) if research_content else "(논문 내용 없음)",
    })

def generate_proposal_draft(llm, project_summary: str, template_guidance: str, research_evidence: str):
    chain = ChatPromptTemplate.from_template(get_proposal_prompt()) | llm | StrOutputParser()
    return chain.invoke({
        "project_summary": truncate_text(project_summary, 12000),
        "template_guidance": truncate_text(template_guidance, 15000),
        "research_evidence": truncate_text(research_evidence, 16000)
    })

def revise_for_compliance(llm, draft: str, template_guidance: str):
    chain = ChatPromptTemplate.from_template(get_compliance_revision_prompt()) | llm | StrOutputParser()
    return chain.invoke({
        "template_guidance": truncate_text(template_guidance, 12000),
        "draft": truncate_text(draft, 42000),
    })


# ==========================================
# ==== 8. Sanitization & Formatting ========
# ==========================================
def regex_sanitize_final_text(text: str) -> str:
    out = text or ""
    
    # <br> 태그를 실제 줄바꿈으로 변경
    out = out.replace("<br>", "\n").replace("&lt;br&gt;", "\n").replace("<br/>", "\n")
    
    # 1인칭 및 메타 표현 제거
    direct_replacements = {
        "본 제안서는": "", "필자": "", 
        "논문에서는": "동국대학교 연구진의 선행연구에서는",
        "논문에서": "동국대학교 선행연구에서",
    }
    for k, v in direct_replacements.items():
        out = out.replace(k, v)

    # 임의 태그 강제 삭제
    out = re.sub(r"\[\[TABLE:[^\]]+\]\]", "", out)
    out = re.sub(r"\[\[IMAGE:[A-Z_]+(?<!GOAL)(?<!FRAMEWORK)\]\]", "", out)
    
    # 잔여 마크다운 아티팩트 제거
    out = out.replace("**", "")
    out = re.sub(r"\n0\s*\n", "\n", out)
    out = re.sub(r"\n{3,}", "\n\n", out)
    
    return out.strip()


# ==========================================
# ==== 9. DOCX Export ======================
# ==========================================
def set_doc_default_style(doc, font_name="바탕", font_size=10.5):
    style = doc.styles["Normal"]
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(font_size)

def add_doc_paragraph(doc, text, bold=False, size=10.5, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "바탕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "바탕")
    run.font.size = Pt(size)
    return p

def _parse_markdown_table(lines: List[str]) -> Optional[List[List[str]]]:
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            break
        if re.match(r"^\|[\s\-:|]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    return rows if rows else None

def _add_docx_table(doc, rows: List[List[str]]):
    if not rows: return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci < num_cols:
                cell = table.rows[ri].cells[ci]
                cell.text = cell_text
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = "바탕"
                        run.font.size = Pt(9)
                        if ri == 0:
                            run.bold = True

def _add_docx_markdown_block(doc, block: str):
    block = block.strip()
    if not block: return

    if block.startswith("# "):
        add_doc_paragraph(doc, block[2:].strip(), bold=True, size=14)
        return
    if block.startswith("## "):
        add_doc_paragraph(doc, block[3:].strip(), bold=True, size=12.5)
        return
    if block.startswith("### "):
        add_doc_paragraph(doc, block[4:].strip(), bold=True, size=11.5)
        return

    lines = block.splitlines()
    if lines and lines[0].strip().startswith("|"):
        rows = _parse_markdown_table(lines)
        if rows:
            _add_docx_table(doc, rows)
            return

    if all(line.strip().startswith("- ") for line in lines if line.strip()):
        for line in lines:
            if not line.strip(): continue
            p = doc.add_paragraph(style=None)
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run("• " + line.strip()[2:].strip())
            run.font.name = "바탕"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "바탕")
            run.font.size = Pt(10.5)
        return

    for line in lines:
        if not line.strip():
            doc.add_paragraph()
            continue
        p = doc.add_paragraph()
        run = p.add_run(line.strip())
        run.font.name = "바탕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "바탕")
        run.font.size = Pt(10.5)

def build_docx_bytes(md_text: str, goal_img=None, fw_img=None) -> Optional[bytes]:
    if DocxDocument is None:
        return None
    doc = DocxDocument()
    set_doc_default_style(doc, font_name="바탕", font_size=10.5)
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    tokens = re.split(r"(\[\[IMAGE:GOAL\]\]|\[\[IMAGE:FRAMEWORK\]\])", md_text)
    
    for token in tokens:
        if not token:
            continue
        
        if token == "[[IMAGE:GOAL]]":
            if goal_img is not None:
                buf = BytesIO()
                goal_img.save(buf, format="PNG")
                buf.seek(0)
                doc.add_picture(buf, width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                add_doc_paragraph(doc, "[목표·비전 개념도 삽입 위치]", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            continue
        
        if token == "[[IMAGE:FRAMEWORK]]":
            if fw_img is not None:
                buf = BytesIO()
                fw_img.save(buf, format="PNG")
                buf.seek(0)
                doc.add_picture(buf, width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                add_doc_paragraph(doc, "[연구 프레임워크 도식 삽입 위치]", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            continue
            
        blocks = token.split("\n\n")
        for block in blocks:
            _add_docx_markdown_block(doc, block)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


# ==========================================
# ==== 10. UI Layout & Main Logic ==========
# ==========================================

# 헤더 (로고 제거, 제목만 출력)
st.markdown("""
<div class="header-container">
    <h1 class="header-title">사업계획서 초안 작성기</h1>
</div>
""", unsafe_allow_html=True)

# 사이드바
with st.sidebar:
    st.header("API 설정")
    api_key_input = st.text_input("Google AI Studio API Key", type="password", value=DEFAULT_API_KEY)
    if api_key_input:
        os.environ["GOOGLE_API_KEY"] = api_key_input
    st.divider()

# 메인 파일 업로드 영역
col1, col2 = st.columns(2)

with col1:
    st.markdown("### 공고 및 양식 파일")
    notice_pdf_upload = st.file_uploader("1. 사업 공고문 (PDF)", type=["pdf"], key="notice_pdf")
    template_pdf_upload = st.file_uploader("2. 제안서 양식 (PDF)", type=["pdf"], key="template_pdf")

with col2:
    st.markdown("### 논문 파일")
    paper_pdf_uploads = st.file_uploader("교수님 논문 (PDF)", type=["pdf"], accept_multiple_files=True, key="paper_pdfs")

st.markdown("<br>", unsafe_allow_html=True)

# 🚀 생성 버튼
generate_btn = st.button("과제 요약 및 제안서 생성", use_container_width=True)
st.markdown("---")

# 생성 로직
if generate_btn:
    if not os.environ.get("GOOGLE_API_KEY"):
        st.error("사이드바에 Google API Key를 입력하세요.")
        st.stop()
    if not notice_pdf_upload:
        st.error("사업 공고문 PDF를 업로드해 주세요.")
        st.stop()
    if not paper_pdf_uploads:
        st.error("논문 PDF를 최소 1개 이상 업로드해 주세요.")
        st.stop()

    api_key = os.environ["GOOGLE_API_KEY"]

    with st.status("데이터 분석 및 제안서 생성 중...", expanded=True) as status:
        
        # 1. 공고 요약 (FAISS + Flash)
        st.write("1️⃣ 공고문을 분석하여 핵심 요약을 생성합니다...")
        notice_docs = extract_docs_from_uploaded_pdf(notice_pdf_upload)
        summary_result = run_summarization(notice_docs, api_key)
        st.session_state.summary_result = summary_result
        
        # 2. 공고 및 양식 텍스트 추출
        st.write("2️⃣ 공고 및 양식의 텍스트와 표 구조를 추출합니다...")
        notice_text, _ = extract_text_from_uploaded_pdf(notice_pdf_upload)
        template_text, _ = extract_text_from_uploaded_pdf(template_pdf_upload) if template_pdf_upload else ("", "")
        
        llm_pro = build_text_llm(api_key, FIXED_MODEL_PRO)
        template_guide = analyze_notice_and_template(llm_pro, notice_text, template_text)
        st.session_state.template_guide = template_guide

        # 3. 논문 텍스트 추출
        st.write(f"3️⃣ 업로드된 논문 {len(paper_pdf_uploads)}건의 기술 요소를 추출합니다...")
        combined_paper_text = ""
        for paper_file in paper_pdf_uploads:
            text, name = extract_text_from_uploaded_pdf(paper_file, max_pages=12) 
            combined_paper_text += f"\n\n--- [논문: {name}] ---\n{text}"
        
        research_ev = build_research_evidence(llm_pro, combined_paper_text)
        st.session_state.research_ev = research_ev

        # 4. 제안서 초안 생성
        st.write("4️⃣ 동국대학교 맞춤형 제안서 초안을 작성합니다...")
        draft = generate_proposal_draft(
            llm=llm_pro,
            project_summary=summary_result,
            template_guidance=template_guide,
            research_evidence=research_ev
        )

        # 5. 검토 및 후처리
        st.write("5️⃣ 서식 준수 여부를 검토하고 텍스트를 정제합니다...")
        revised = revise_for_compliance(llm_pro, draft, template_guide)
        revised = regex_sanitize_final_text(revised)
        st.session_state.proposal_result = revised
        
        # 6. 이미지 생성 (목표 및 프레임워크 도식)
        st.write("6️⃣ 도식(목표, 프레임워크) 이미지를 생성합니다...")
        genai_client = build_genai_client(api_key)
        
        goal_text, goal_img = generate_goal_image_with_genai(
            client=genai_client,
            project_summary=summary_result,
            template_guidance=template_guide,
            research_evidence=research_ev
        )
        st.session_state.goal_img = goal_img
        
        fw_text, fw_img = generate_framework_image_with_genai(
            client=genai_client,
            project_summary=summary_result,
            template_guidance=template_guide,
            research_evidence=research_ev
        )
        st.session_state.fw_img = fw_img

        # 7. DOCX 바이트 생성 후 세션 저장
        st.write("7️⃣ Word(DOCX) 파일을 생성합니다...")
        st.session_state.docx_bytes = build_docx_bytes(revised, goal_img, fw_img)

        status.update(label="생성 완료!", state="complete", expanded=False)


# ==========================================
# ==== 11. 결과 출력 (세션 상태 유지) ========
# ==========================================
if st.session_state.summary_result:
    st.markdown("### 공고 핵심 요약")
    st.markdown(f'<div class="white-box">{st.session_state.summary_result}</div>', unsafe_allow_html=True)

if st.session_state.proposal_result:
    st.markdown("### 삽입 예정 도식 (생성된 이미지)")
    img_col1, img_col2 = st.columns(2)
    with img_col1:
        if st.session_state.goal_img:
            st.image(st.session_state.goal_img, use_container_width=True, caption="목표·비전 개념도")
        else:
            st.markdown('<div class="image-placeholder"><h4>목표·비전 개념도</h4><p>이미지 생성에 실패했습니다.</p></div>', unsafe_allow_html=True)
    
    with img_col2:
        if st.session_state.fw_img:
            st.image(st.session_state.fw_img, use_container_width=True, caption="연구 프레임워크 도식")
        else:
            st.markdown('<div class="image-placeholder"><h4>연구 프레임워크 도식</h4><p>이미지 생성에 실패했습니다.</p></div>', unsafe_allow_html=True)

    st.markdown("### 동국대학교 제안서 초안")
    display_text = st.session_state.proposal_result.replace("[[IMAGE:GOAL]]", "").replace("[[IMAGE:FRAMEWORK]]", "")
    st.markdown(f'<div class="white-box">{display_text}</div>', unsafe_allow_html=True)

    st.markdown("---")
    st.subheader("결과물 저장")
    st.caption("저장 버튼을 눌러도 화면이 초기화되지 않습니다.")

    btn_col1, btn_col2, btn_col3 = st.columns(3)
    
    with btn_col1:
        md_download = st.session_state.proposal_result.replace("[[IMAGE:GOAL]]", "\n\n*(목표·비전 개념도 위치)*\n\n").replace("[[IMAGE:FRAMEWORK]]", "\n\n*(연구 프레임워크 도식 위치)*\n\n")
        st.download_button(
            label="마크다운(.md) 저장",
            data=md_download,
            file_name="동국대_제안서_초안.md",
            mime="text/markdown",
            use_container_width=True
        )

    with btn_col2:
        if st.session_state.docx_bytes:
            st.download_button(
                label="워드(.docx) 저장",
                data=st.session_state.docx_bytes,
                file_name="동국대_제안서_초안.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

    with btn_col3:
        zip_buf = BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("동국대_제안서_초안.md", md_download)
            if st.session_state.docx_bytes:
                zf.writestr("동국대_제안서_초안.docx", st.session_state.docx_bytes)
            zf.writestr("공고요약.txt", st.session_state.summary_result)
            zf.writestr("논문_기술요소분석.txt", st.session_state.research_ev)
            zf.writestr("양식가이드.txt", st.session_state.template_guide)
            
            # 생성된 이미지가 있다면 ZIP 파일에 포함
            if st.session_state.goal_img:
                buf = BytesIO()
                st.session_state.goal_img.save(buf, format="PNG")
                zf.writestr("목표_개념도.png", buf.getvalue())
            if st.session_state.fw_img:
                buf = BytesIO()
                st.session_state.fw_img.save(buf, format="PNG")
                zf.writestr("연구_프레임워크.png", buf.getvalue())
            
        st.download_button(
            label="전체 파일 ZIP 저장",
            data=zip_buf.getvalue(),
            file_name="동국대_제안서_패키지.zip",
            mime="application/zip",
            use_container_width=True
        )

elif not notice_pdf_upload:
    st.info("좌측 상단에서 사업 공고문 PDF를 업로드해 주세요.")